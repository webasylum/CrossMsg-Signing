#!/usr/bin/env python3

import sys
import os
from docx import Document
import json
from typing import Dict, List, Set, Any
import re

JAVA_RULES_PATH = "src/main/java/com/tsg/crossmsg/signing/config/ConversionRules.java"

RULE_KEYWORDS = [
    'must', 'shall', 'should', 'required', 'constraint', 'pattern', 'encoding', 'array', 'object',
    'minItems', 'maxItems', 'namespace', 'enum', 'description', 'type', 'ref', 'property', 'element',
    'cardinality', 'format', 'identifier', 'name', 'value', 'schema', 'set', 'list', 'map', 'validation',
    'error', 'exception', 'limit', 'length', 'digits', 'decimal', 'date', 'time', 'boolean', 'code',
    'amount', 'currency', 'number', 'string', 'object', 'array', 'anchor', 'additionalProperties',
    'minProperties', 'maxProperties', 'minLength', 'maxLength', 'minInclusive', 'maxInclusive',
    'minExclusive', 'maxExclusive', 'pattern', 'enum', 'description', 'definition', 'reference',
    'component', 'choice', 'message', 'document', 'file', 'urn', 'iso', 'json', 'xml', 'draft', '2020-12',
    'section', '8.', '9.', '10.', '11.'
]

class ConversionRulesValidator:
    def __init__(self):
        # Constants from ConversionRules.java
        self.DESCRIPTION_SEPARATOR = " \n"
        self.ARRAY_TYPE = "array"
        self.STRING_TYPE = "string"
        self.DEFAULT_MIN_ITEMS = 1
        self.JSON_SCHEMA_VERSION = "https://json-schema.org/draft/2020-12/schema"
        self.BASE_URN = "urn:iso:std:iso:20022:tech:json:"
        self.SCHEMA_TYPE = "object"
        self.ADDITIONAL_PROPERTIES = False
        self.ROOT_ELEMENT = "Document"
        self.ENCODING = "UTF-8"
        self.IDENTIFIER_SEPARATOR = "."
        self.SCHEMA_FILE_SUFFIX = ".schema.json"
        self.NAME_SEPARATOR = "-"
        self.DEFAULT_NAMESPACE = "urn:iso:std:iso:20022:tech:xsd:pacs.008.001.09"
        self.BIZ_MSG_ENVLP_NAMESPACE = "urn:iso:std:iso:20022:tech:xsd:head.001.001.02"
        
        # Initialize validation results
        self.validation_results = {
            "total_paragraphs": 0,
            "schema_elements": 0,
            "validation_errors": [],
            "warnings": [],
            "rule_matches": {
                "namespaces": [],
                "schema_types": [],
                "array_definitions": [],
                "required_elements": [],
                "currency_handling": [],
                "encoding_rules": []
            }
        }

    def validate_document(self, doc: Document) -> Dict[str, Any]:
        """Validate the Word document against ConversionRules.java"""
        self.validation_results["total_paragraphs"] = len(doc.paragraphs)
        
        # Process paragraphs in chunks
        chunk_size = 100
        for i in range(0, len(doc.paragraphs), chunk_size):
            chunk = doc.paragraphs[i:i + chunk_size]
            self._process_chunk(chunk)
            
        return self.validation_results

    def _process_chunk(self, paragraphs: List[Any]) -> None:
        """Process a chunk of paragraphs from the document"""
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Check for schema elements
            if "schema" in text.lower():
                self.validation_results["schema_elements"] += 1
                
            # Validate against conversion rules
            self._validate_against_rules(text)

    def _validate_against_rules(self, text: str) -> None:
        """Validate text against conversion rules"""
        # Check namespace rules
        if "namespace" in text.lower():
            if "urn:iso:std:iso:20022" in text:
                self.validation_results["rule_matches"]["namespaces"].append(text[:100])
                
        # Check schema type rules
        if "type" in text.lower() and "object" in text.lower():
            if "schema" in text.lower():
                self.validation_results["rule_matches"]["schema_types"].append(text[:100])
                
        # Check array definitions
        if "array" in text.lower():
            if "minItems" in text.lower() or "maxItems" in text.lower():
                self.validation_results["rule_matches"]["array_definitions"].append(text[:100])
                
        # Check required elements
        if "required" in text.lower():
            if "element" in text.lower():
                self.validation_results["rule_matches"]["required_elements"].append(text[:100])
                
        # Check currency handling
        if "currency" in text.lower() or "ccy" in text.lower():
            if "amount" in text.lower():
                self.validation_results["rule_matches"]["currency_handling"].append(text[:100])
                
        # Check encoding rules
        if "encoding" in text.lower() or "utf-8" in text.lower():
            self.validation_results["rule_matches"]["encoding_rules"].append(text[:100])

def extract_java_rules(java_path: str) -> Set[str]:
    """Extract all rule-like constants, method names, and docstrings from ConversionRules.java"""
    rules = set()
    if not os.path.exists(java_path):
        print(f"Warning: {java_path} not found!")
        return rules
    with open(java_path, encoding="utf-8") as f:
        for line in f:
            # Constants and static fields
            m = re.match(r'\s*public static final [^ ]+ ([A-Z0-9_]+) ?=', line)
            if m:
                rules.add(m.group(1).lower())
            # Method names
            m = re.match(r'\s*public static [^ ]+ ([a-zA-Z0-9_]+)\(', line)
            if m:
                rules.add(m.group(1).lower())
            # Docstring section references
            m = re.search(r'Section ([0-9.]+)', line)
            if m:
                rules.add(m.group(0).lower())
            # Any quoted string that looks like a rule
            for quoted in re.findall(r'"([^"]+)"', line):
                if any(kw in quoted.lower() for kw in RULE_KEYWORDS):
                    rules.add(quoted.strip().lower())
    return rules

def extract_word_rules(doc: Document) -> List[str]:
    """Extract rule-like statements from the Word document"""
    rules = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if any(kw in text.lower() for kw in RULE_KEYWORDS):
            rules.append(text)
    return rules

def main():
    doc_path = os.path.join('docs', 'iso', 'ISO 20022 - JSON Schema Draft 2020-12 generation (v20250321) (clean).docx')
    java_path = JAVA_RULES_PATH
    
    if not os.path.exists(doc_path):
        print(f"Error: Document not found at {doc_path}")
        sys.exit(1)
    if not os.path.exists(java_path):
        print(f"Error: ConversionRules.java not found at {java_path}")
        sys.exit(1)
    
    print("Extracting rules from ConversionRules.java...")
    java_rules = extract_java_rules(java_path)
    print(f"Found {len(java_rules)} rule-like entries in ConversionRules.java.")
    
    print("\nExtracting rule-like statements from the Word document...")
    doc = Document(doc_path)
    word_rules = extract_word_rules(doc)
    print(f"Found {len(word_rules)} rule-like statements in the Word document.")
    
    print("\nComparing rules...")
    unmatched = []
    for rule in word_rules:
        # Check if any substring of the rule is in java_rules
        found = False
        rule_lower = rule.lower()
        for jr in java_rules:
            if jr in rule_lower or rule_lower in jr:
                found = True
                break
        if not found:
            unmatched.append(rule)
    
    print(f"\nRules in the Word document NOT found in ConversionRules.java:")
    if unmatched:
        for rule in unmatched:
            print(f"- {rule}")
    else:
        print("All rule-like statements from the Word document are represented in ConversionRules.java.")

if __name__ == "__main__":
    main() 